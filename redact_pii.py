#!/usr/bin/env python3
"""Create a de-identified DOCX copy while retaining its paragraph and table structure.

The detector is deliberately conservative for free-form names: named entities are
redacted when they appear in an identifying context (for example, ``Contact
Person`` or ``Director``), while email, phone, IP, payment, DOB, address and
legal-entity patterns are detected everywhere.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
from collections import Counter
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Iterator, Mapping

from docx import Document
from docx.text.paragraph import Paragraph


@dataclass(frozen=True)
class Span:
    start: int
    end: int
    pii_type: str
    value: str


EMAIL_RE = re.compile(r"(?i)\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b")
PHONE_RE = re.compile(
    r"(?<!\w)(?:\+?\s?\d{1,3}[\s.-]?)?(?:\(?0?\d{2,5}\)?[\s.-]?)?\d{3,5}[\s.-]\d{4,6}(?!\w)"
)
PHONE_CONTEXT_RE = re.compile(
    r"(?i)\b(?:phone|mobile|cell|telephone|tel|contact\s+number|contact\s+no\.?|whatsapp)\s*[:#-]?\s*(\+?\d[\d\s().-]{8,}\d)"
)
SSN_RE = re.compile(r"(?<!\d)\d{3}-\d{2}-\d{4}(?!\d)")
IP_RE = re.compile(r"(?<![\d.])(?:25[0-5]|2[0-4]\d|1?\d?\d)(?:\.(?:25[0-5]|2[0-4]\d|1?\d?\d)){3}(?![\d.])")
DATE_RE = re.compile(
    r"\b(?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}[/-]\d{1,2}[/-]\d{1,2}|\d{1,2}\s+(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Sept|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{4})\b",
    re.I,
)
CARD_CANDIDATE_RE = re.compile(r"(?<!\d)(?:\d[ -]?){13,19}\d(?!\d)")
PIN_RE = re.compile(r"\b\d{3}\s?\d{3}\b")
POSTAL_RE = re.compile(r"\b(?:\d{3}\s?\d{3}|\d{5}(?:-\d{4})?)\b")
ORG_RE = re.compile(
    r"\b(?:[A-Z][A-Za-z0-9&.'-]*,?\s+){1,8}(?:Private\s+Limited|Public\s+Limited|Pvt\.?\s+Ltd\.?|Limited|Ltd\.?|LLP|L\.L\.P\.|LLC|L\.L\.C\.|Inc\.?|Incorporated|Corporation|Corp\.?|Company|Co\.?|Bank|PLC|GmbH)\b"
)
NAME_CONTEXT_RE = re.compile(
    r"(?i)\b(?:contact\s+person|full\s+name|customer\s+name|client\s+name|applicant\s+name|candidate\s+name|employee\s+name|patient\s+name|student\s+name|user\s+name|name|director|managing\s+director|chief\s+financial\s+officer|company\s+secretary|compliance\s+officer|partner|promoter|auditor)\s*:\s*"
    r"((?:Mr\.|Ms\.|Mrs\.|Dr\.)?\s*[A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})"
)
TITLED_NAME_RE = re.compile(r"\b(?:Mr\.|Ms\.|Mrs\.|Dr\.)\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3}\b")
ADDRESS_LABEL_RE = re.compile(
    r"(?i)\b(?:registered\s+office|corporate\s+office|home\s+address|office\s+address|billing\s+address|shipping\s+address|residence|mailing\s+address|address)\s*:\s*([^\n]+)"
)
ADDRESS_HINT_RE = re.compile(
    r"(?i)\b(?:road|rd\.?|street|st\.?|avenue|ave\.?|lane|ln\.?|drive|dr\.?|boulevard|blvd\.?|suite|apt\.?|apartment|floor|building|tower|office|po\s+box|village|taluka|nagar|mumbai|pune|india|usa|united\s+states)\b"
)
ADDRESS_LINE_RE = re.compile(r"(?im)^.*(?:\b\d{3}\s?\d{3}\b|\b\d{5}(?:-\d{4})?\b).*$")
CONTACT_VALUE_RE = re.compile(r"(?i)\bcontact\s+person\s*:\s*([^;\n]+)")
PERSON_IN_VALUE_RE = re.compile(r"(?:Mr\.|Ms\.|Mrs\.|Dr\.)?\s*[A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3}")
LABEL_TYPES = {
    "name": ("name", "full name", "customer name", "client name", "applicant name", "candidate name", "employee name", "patient name", "student name", "contact person"),
    "dob": ("dob", "date of birth", "birth date", "birthdate"),
    "phone": ("phone", "mobile", "telephone", "tel", "contact number", "contact no", "whatsapp"),
    "address": ("address", "home address", "office address", "billing address", "shipping address", "residence", "mailing address", "registered office"),
    "company": ("company", "employer", "organisation", "organization"),
}


def _luhn(number: str) -> bool:
    digits = [int(c) for c in re.sub(r"\D", "", number)]
    if not 13 <= len(digits) <= 19:
        return False
    total = 0
    for index, digit in enumerate(reversed(digits)):
        if index % 2:
            digit *= 2
            if digit > 9:
                digit -= 9
        total += digit
    return total % 10 == 0


def _overlap(a: Span, b: Span) -> bool:
    return a.start < b.end and b.start < a.end


def _label_type(text: str) -> str | None:
    normalized = re.sub(r"[^a-z0-9]+", " ", text.casefold()).strip()
    if not normalized or len(normalized) > 45:
        return None
    for pii_type, labels in LABEL_TYPES.items():
        if normalized in labels:
            return pii_type
    return None


class Pseudonymizer:
    """Creates deterministic, internally consistent fake values per source value."""

    FIRST = ("Aarav", "Maya", "Dev", "Nina", "Arjun", "Isha", "Rohan", "Tara")
    LAST = ("Shah", "Mehta", "Kapoor", "Rao", "Singh", "Patel", "Gupta", "Verma")

    def __init__(self) -> None:
        self.replacements: dict[tuple[str, str], str] = {}
        self.counters: Counter[str] = Counter()

    @staticmethod
    def _key(value: str) -> str:
        return re.sub(r"\s+", " ", value.strip()).casefold()

    def replacement(self, pii_type: str, value: str) -> str:
        key = (pii_type, self._key(value))
        if key in self.replacements:
            return self.replacements[key]
        self.counters[pii_type] += 1
        ordinal = self.counters[pii_type]
        digest = int(hashlib.sha256(key[1].encode("utf-8")).hexdigest()[:8], 16)
        if pii_type == "name":
            fake = f"{self.FIRST[digest % len(self.FIRST)]} {self.LAST[(digest // 11) % len(self.LAST)]}"
        elif pii_type == "email":
            fake = f"contact{ordinal:03d}@example.com"
        elif pii_type == "phone":
            fake = f"+91 90000 {ordinal:05d}"
        elif pii_type == "company":
            fake = f"Example Entity {ordinal:03d} Limited"
        elif pii_type == "address":
            fake = f"{100 + ordinal} Example Avenue, Sample City - 400{ordinal % 1000:03d}, India"
        elif pii_type == "ssn":
            fake = f"900-{ordinal % 100:02d}-{1000 + ordinal:04d}"
        elif pii_type == "card":
            fake = self._fake_card(ordinal)
        elif pii_type == "dob":
            fake = f"01/01/{1980 + ordinal % 30}"
        elif pii_type == "ip":
            fake = f"203.0.113.{(ordinal % 250) + 1}"
        else:
            raise ValueError(f"Unsupported PII type: {pii_type}")
        self.replacements[key] = fake
        return fake

    @staticmethod
    def _fake_card(ordinal: int) -> str:
        base = f"411111111111{ordinal % 10000:04d}"[:15]
        total = 0
        for index, char in enumerate(reversed(base)):
            digit = int(char)
            if index % 2 == 0:
                digit *= 2
                if digit > 9:
                    digit -= 9
            total += digit
        check = (10 - total % 10) % 10
        number = base + str(check)
        return " ".join(number[i : i + 4] for i in range(0, len(number), 4))


def _add(spans: list[Span], start: int, end: int, pii_type: str, text: str) -> None:
    value = text[start:end]
    if value.strip():
        spans.append(Span(start, end, pii_type, value))


def detect_pii(
    text: str,
    known_names: Iterable[str] = (),
    known_companies: Iterable[str] = (),
    known_values: Mapping[str, Iterable[str]] | None = None,
) -> list[Span]:
    """Return non-overlapping spans, prioritising longer/contextual detections."""
    candidates: list[Span] = []

    # Addresses are intentionally contextual to avoid treating ordinary prose as an address.
    for match in ADDRESS_LABEL_RE.finditer(text):
        address = match.group(1).strip()
        begin = match.start(1) + (len(match.group(1)) - len(match.group(1).lstrip()))
        if POSTAL_RE.search(address) or ADDRESS_HINT_RE.search(address):
            _add(candidates, begin, match.end(1), "address", text)
    for match in ADDRESS_LINE_RE.finditer(text):
        line = match.group(0)
        if ADDRESS_HINT_RE.search(line):
            _add(candidates, match.start(), match.end(), "address", text)

    for match in EMAIL_RE.finditer(text):
        _add(candidates, match.start(), match.end(), "email", text)
    for match in SSN_RE.finditer(text):
        _add(candidates, match.start(), match.end(), "ssn", text)
    for match in IP_RE.finditer(text):
        _add(candidates, match.start(), match.end(), "ip", text)
    for match in CARD_CANDIDATE_RE.finditer(text):
        if _luhn(match.group(0)):
            _add(candidates, match.start(), match.end(), "card", text)
    for match in PHONE_RE.finditer(text):
        raw = match.group(0)
        digits = re.sub(r"\D", "", raw)
        if len(digits) >= 10 and ("+" in raw or " " in raw or "-" in raw or "(" in raw):
            _add(candidates, match.start(), match.end(), "phone", text)
    for match in PHONE_CONTEXT_RE.finditer(text):
        raw = match.group(1)
        digits = re.sub(r"\D", "", raw)
        if 10 <= len(digits) <= 15:
            _add(candidates, match.start(1), match.end(1), "phone", text)
    for match in DATE_RE.finditer(text):
        context = text[max(0, match.start() - 45) : match.end() + 10].casefold()
        if any(marker in context for marker in ("date of birth", "dob", "born", "birth date", "birthdate")):
            _add(candidates, match.start(), match.end(), "dob", text)
    for match in NAME_CONTEXT_RE.finditer(text):
        _add(candidates, match.start(1), match.end(1), "name", text)
    for match in CONTACT_VALUE_RE.finditer(text):
        for person in PERSON_IN_VALUE_RE.finditer(match.group(1)):
            _add(candidates, match.start(1) + person.start(), match.start(1) + person.end(), "name", text)
    for match in TITLED_NAME_RE.finditer(text):
        _add(candidates, match.start(), match.end(), "name", text)
    for match in ORG_RE.finditer(text):
        candidate = match.group(0).strip()
        if len(candidate.split()) <= 9 and not candidate.casefold().startswith("the "):
            _add(candidates, match.start(), match.end(), "company", text)
    for name in known_names:
        for match in re.finditer(r"(?<!\w)" + re.escape(name) + r"(?!\w)", text):
            _add(candidates, match.start(), match.end(), "name", text)
    for company in known_companies:
        for match in re.finditer(r"(?<!\w)" + re.escape(company) + r"(?!\w)", text):
            _add(candidates, match.start(), match.end(), "company", text)
    for pii_type, values in (known_values or {}).items():
        for value in values:
            if not value.strip():
                continue
            for match in re.finditer(r"(?<!\w)" + re.escape(value.strip()) + r"(?!\w)", text):
                _add(candidates, match.start(), match.end(), pii_type, text)

    # De-duplicate and resolve overlaps: contextual/long spans first, then earliest.
    priority = {"address": 0, "email": 1, "ssn": 1, "card": 1, "phone": 1, "ip": 1, "dob": 1, "name": 2, "company": 3}
    chosen: list[Span] = []
    for span in sorted(candidates, key=lambda s: (priority[s.pii_type], -(s.end - s.start), s.start)):
        if not any(_overlap(span, existing) for existing in chosen):
            chosen.append(span)
    return sorted(chosen, key=lambda s: s.start)


def _copy_run_properties(source, target) -> None:
    if source._r.rPr is not None:
        target._r.get_or_add_rPr().append(deepcopy(source._r.rPr))


def replace_paragraph(
    paragraph: Paragraph,
    pseudonymizer: Pseudonymizer,
    counts: Counter[str],
    known_names: Iterable[str],
    known_companies: Iterable[str],
    known_values: Mapping[str, Iterable[str]] | None = None,
) -> bool:
    original = paragraph.text
    spans = detect_pii(original, known_names, known_companies, known_values)
    if not spans:
        return False
    runs = list(paragraph.runs)
    if not runs:
        return False
    boundaries: list[tuple[int, int, object]] = []
    offset = 0
    for run in runs:
        boundaries.append((offset, offset + len(run.text), run))
        offset += len(run.text)

    pieces: list[tuple[str, object]] = []
    cursor = 0
    for span in spans:
        for start, end, run in boundaries:
            left, right = max(cursor, start), min(span.start, end)
            if left < right:
                pieces.append((original[left:right], run))
        style_run = next((run for start, end, run in boundaries if start <= span.start < end), runs[0])
        pieces.append((pseudonymizer.replacement(span.pii_type, span.value), style_run))
        counts[span.pii_type] += 1
        cursor = span.end
    for start, end, run in boundaries:
        left, right = max(cursor, start), min(len(original), end)
        if left < right:
            pieces.append((original[left:right], run))

    paragraph._p.clear_content()
    for value, source_run in pieces:
        if value:
            new_run = paragraph.add_run(value)
            _copy_run_properties(source_run, new_run)
    return True


def iter_paragraphs(document: Document) -> Iterator[Paragraph]:
    seen: set[int] = set()

    def emit(paragraph: Paragraph) -> Iterator[Paragraph]:
        key = id(paragraph._p)
        if key not in seen:
            seen.add(key)
            yield paragraph

    def walk_table(table) -> Iterator[Paragraph]:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield from emit(paragraph)
                for nested in cell.tables:
                    yield from walk_table(nested)

    for paragraph in document.paragraphs:
        yield from emit(paragraph)
    for table in document.tables:
        yield from walk_table(table)
    for section in document.sections:
        for story in (section.header, section.footer, section.first_page_header, section.first_page_footer):
            for paragraph in story.paragraphs:
                yield from emit(paragraph)
            for table in story.tables:
                yield from walk_table(table)


def _contextual_known_values(paragraphs: list[Paragraph]) -> dict[str, set[str]]:
    """Learn PII values from adjacent label/value paragraphs.

    DOCX tables often expose each cell as a separate paragraph, so a row such as
    `Full Name | Marcus Hill` becomes two neighbouring paragraphs. This pass
    lets the redactor work on arbitrary forms/tickets, not just prose where the
    label and value share one paragraph.
    """
    known: dict[str, set[str]] = {kind: set() for kind in LABEL_TYPES}
    for previous, current in zip(paragraphs, paragraphs[1:]):
        label = _label_type(previous.text)
        value = current.text.strip()
        if not label or not value or len(value) > 180:
            continue
        if label == "name" and PERSON_IN_VALUE_RE.fullmatch(value):
            known["name"].add(value)
        elif label == "dob":
            for match in DATE_RE.finditer(value):
                known["dob"].add(match.group(0))
        elif label == "phone":
            digits = re.sub(r"\D", "", value)
            if 10 <= len(digits) <= 15:
                known["phone"].add(value)
        elif label == "address" and (POSTAL_RE.search(value) or ADDRESS_HINT_RE.search(value)):
            known["address"].add(value)
        elif label == "company":
            if ORG_RE.search(value) or len(value.split()) >= 2:
                known["company"].add(value)
    return {kind: values for kind, values in known.items() if values}


def redact_docx(source: Path, output: Path, mapping_path: Path | None = None) -> dict:
    document = Document(source)
    pseudonymizer = Pseudonymizer()
    counts: Counter[str] = Counter()
    changed_paragraphs = 0
    paragraphs = list(iter_paragraphs(document))
    seed_spans = [span for paragraph in paragraphs for span in detect_pii(paragraph.text)]
    known_values = _contextual_known_values(paragraphs)
    known_names = sorted(
        {span.value.strip() for span in seed_spans if span.pii_type == "name"} | set(known_values.get("name", set())),
        key=len,
        reverse=True,
    )
    known_companies = sorted(
        {span.value.strip() for span in seed_spans if span.pii_type == "company" and len(span.value.split()) > 1}
        | set(known_values.get("company", set())),
        key=len,
        reverse=True,
    )
    for paragraph in paragraphs:
        if replace_paragraph(paragraph, pseudonymizer, counts, known_names, known_companies, known_values):
            changed_paragraphs += 1
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    summary = {
        "source": str(source),
        "output": str(output),
        "changed_paragraphs": changed_paragraphs,
        "redactions": dict(sorted(counts.items())),
        "unique_replacements": len(pseudonymizer.replacements),
    }
    if mapping_path:
        # The mapping is intentionally optional; do not submit it with a redacted document.
        mapping_path.write_text(
            json.dumps({f"{kind}:{value}": fake for (kind, value), fake in pseudonymizer.replacements.items()}, indent=2),
            encoding="utf-8",
        )
    return summary


EVALUATION_CASES = [
    ("Contact Person: Rashi Patil", {"name"}),
    ("Mr. Rohan Dey attended the meeting.", {"name"}),
    ("Email rashhi.patil@gmail.com for support.", {"email"}),
    ("Telephone: +91 98765 43210", {"phone"}),
    ("KSH International Limited", {"company"}),
    ("Registered Office: 11/3 Example Road, Pune - 411 001, India", {"address"}),
    ("SSN 123-45-6789", {"ssn"}),
    ("Card: 4111 1111 1111 1111", {"card"}),
    ("Date of Birth: 12/03/1990", {"dob"}),
    ("Login from 192.168.1.10", {"ip"}),
    ("The offer closes on 10 December 2025.", set()),
    ("Corporate Identity Number: U28129PN1979PLC141032", set()),
    ("Order 2025-10001 is valid.", set()),
    ("The Book Running Lead Managers are appointed.", set()),
]


def evaluate() -> dict:
    tp = fp = fn = tn = 0
    per_type: dict[str, Counter[str]] = {kind: Counter() for kind in ("name", "email", "phone", "company", "address", "ssn", "card", "dob", "ip")}
    for text, expected in EVALUATION_CASES:
        found = {span.pii_type for span in detect_pii(text)}
        types = expected | found
        for kind in types:
            if kind in expected and kind in found:
                tp += 1; per_type[kind]["tp"] += 1
            elif kind in found:
                fp += 1; per_type[kind]["fp"] += 1
            else:
                fn += 1; per_type[kind]["fn"] += 1
        if not expected and not found:
            tn += 1
    precision = tp / (tp + fp) if tp + fp else 0.0
    recall = tp / (tp + fn) if tp + fn else 0.0
    accuracy = (tp + tn) / (tp + tn + fp + fn) if tp + tn + fp + fn else 0.0
    return {
        "cases": len(EVALUATION_CASES), "tp": tp, "fp": fp, "fn": fn, "tn": tn,
        "precision": precision, "recall": recall, "accuracy": accuracy,
        "per_type": {kind: dict(values) for kind, values in per_type.items()},
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="Redact PII in a DOCX with deterministic fake replacements.")
    parser.add_argument("input", nargs="?", type=Path, help="Source DOCX")
    parser.add_argument("output", nargs="?", type=Path, help="Redacted DOCX")
    parser.add_argument("--mapping", type=Path, help="Optional local-only JSON mapping (do not distribute).")
    parser.add_argument("--evaluate", action="store_true", help="Print metrics for the deterministic labelled test suite.")
    args = parser.parse_args()
    if args.evaluate:
        print(json.dumps(evaluate(), indent=2))
        return
    if not args.input or not args.output:
        parser.error("input and output are required unless --evaluate is used")
    print(json.dumps(redact_docx(args.input, args.output, args.mapping), indent=2))


if __name__ == "__main__":
    main()
