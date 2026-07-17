"""Verify that hybrid ML/NER mode adds recall beyond rules-only mode.

This test uses unlabelled prose names and organisations that the conservative
rules engine should avoid, then verifies spaCy-backed hybrid mode redacts them.

Run:
    python ml_ner_test.py
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from redact_pii import redact_docx


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "examples" / "ml_ner"
INPUT = OUT_DIR / "unlabelled_entities.docx"
RULES_OUTPUT = OUT_DIR / "unlabelled_entities_rules.docx"
HYBRID_OUTPUT = OUT_DIR / "unlabelled_entities_hybrid.docx"
REPORT = OUT_DIR / "ml_ner_report.json"

UNLABELLED_VALUES = ("Alice Johnson", "Robert Chen", "Microsoft")
STRUCTURED_VALUES = ("alice.johnson@example.com", "+1 415-555-0134")
CONTROL_VALUES = ("Ticket ID CASE-2026-771",)


def create_fixture() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading("Unlabelled Entity NER Fixture", level=1)
    document.add_paragraph("Alice Johnson escalated the account to Robert Chen at Microsoft yesterday.")
    document.add_paragraph("Follow-up email: alice.johnson@example.com")
    document.add_paragraph("Phone: +1 415-555-0134")
    document.add_paragraph("Control value should remain: Ticket ID CASE-2026-771")
    document.save(INPUT)


def document_text(path: Path) -> str:
    document = Document(path)
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def main() -> None:
    create_fixture()
    rules_summary = redact_docx(INPUT, RULES_OUTPUT, mode="rules")
    hybrid_summary = redact_docx(INPUT, HYBRID_OUTPUT, mode="hybrid")
    rules_text = document_text(RULES_OUTPUT)
    hybrid_text = document_text(HYBRID_OUTPUT)
    report = {
        "passed": True,
        "input": str(INPUT),
        "rules_output": str(RULES_OUTPUT),
        "hybrid_output": str(HYBRID_OUTPUT),
        "rules_summary": rules_summary,
        "hybrid_summary": hybrid_summary,
        "rules_unlabelled_values_remaining": [value for value in UNLABELLED_VALUES if value in rules_text],
        "hybrid_unlabelled_values_remaining": [value for value in UNLABELLED_VALUES if value in hybrid_text],
        "hybrid_structured_values_remaining": [value for value in STRUCTURED_VALUES if value in hybrid_text],
        "hybrid_missing_controls": [value for value in CONTROL_VALUES if value not in hybrid_text],
        "hybrid_added_redactions": sum(hybrid_summary["redactions"].values()) - sum(rules_summary["redactions"].values()),
    }
    report["passed"] = (
        set(report["rules_unlabelled_values_remaining"]) == set(UNLABELLED_VALUES)
        and not report["hybrid_unlabelled_values_remaining"]
        and not report["hybrid_structured_values_remaining"]
        and not report["hybrid_missing_controls"]
        and report["hybrid_added_redactions"] >= len(UNLABELLED_VALUES)
    )
    REPORT.write_text(json.dumps(report, indent=2), encoding="utf-8")
    print(json.dumps(report, indent=2))
    if not report["passed"]:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
