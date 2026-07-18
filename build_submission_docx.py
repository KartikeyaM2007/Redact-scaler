from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "assets"
OUT = ROOT / "SUBMISSION.docx"


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def h(text: str, level: int = 1) -> None:
        doc.add_heading(text, level=level)

    def p(text: str) -> None:
        doc.add_paragraph(text)

    def bullet(text: str) -> None:
        doc.add_paragraph(text, style="List Bullet")

    h("PII Redaction Tool — my submission notes")
    p("Kartikeya Mishra")
    p("GitHub: https://github.com/KartikeyaM2007/Redact-scaler")
    p("Try Rules online: https://huggingface.co/spaces/Kartikeym2007/Redact")
    p(
        "I'm submitting this for the Scaler AI Labs PII redaction task. "
        "Numbers below were re-checked with python verify_metrics.py on 18 Jul 2026 "
        "(examples/verified_metrics.json)."
    )

    h("The job, in my words", 2)
    p(
        "I needed a script that reads a Word file full of personal data "
        "(they gave a Red Herring Prospectus; the same code also works on ticket-style docs) "
        "and writes a second .docx where real PII is swapped for fake but believable values. "
        "Masking with **** wasn't the ask — they wanted stand-ins like a fake name/email "
        "so the doc still looks readable."
    )
    p("I cover at least: names, emails, phones, companies, addresses, SSNs, cards (Luhn), DOBs, IPs.")
    p("Core file: redact_pii.py")
    p("My redacted prospectus: Red Herring Prospectus - Redacted.docx")
    p("Optional UI: web_app.py")
    p("I do not redact order/ticket/CIN-style IDs. That's a deliberate precision choice.")

    h("How detection works (two modes)", 2)
    p(
        "Rules — regexes + labels I care about (Contact Person, DOB, Registered Office, "
        "Ltd/LLC endings, label→value table cells). Good for structured PII; I can usually explain each hit."
    )
    p(
        "Hybrid (ML / NER) — rules + spaCy en_core_web_sm. Added because rules ignore bare "
        "names/companies in prose. Verified: Rules = 2 hits (email+phone); Hybrid = 5 "
        "(adds Alice Johnson, Robert Chen, Microsoft). See ml_ner_test.py."
    )

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Mode"
    table.cell(0, 1).text = "ml_ner_test.py (live)"
    table.cell(1, 0).text = "Rules"
    table.cell(1, 1).text = "2 redactions (email + phone)"
    table.cell(2, 0).text = "Hybrid"
    table.cell(2, 1).text = "5 redactions (+ 2 names + Microsoft)"

    h("Screens from my local UI", 3)
    p("Rules:")
    doc.add_picture(str(ASSETS / "frontend-rules-mode.png"), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p("Hybrid:")
    doc.add_picture(str(ASSETS / "frontend-ml-ner-mode.png"), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    h("Live demo vs local", 2)
    p("HF Space is static → no Python spaCy → ML toggle stays off on purpose.")
    doc.add_picture(str(ASSETS / "hf-space-ml-disabled.png"), width=Inches(4.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p("Local: pip install -r requirements.txt then python web_app.py")

    h("Verified evaluation numbers", 2)
    h("Accuracy / precision / recall (redact_pii.py --evaluate)", 3)
    t3 = doc.add_table(rows=9, cols=2)
    t3.style = "Table Grid"
    for i, (a, b) in enumerate(
        [
            ("Metric", "Value"),
            ("Cases", "14"),
            ("TP", "10"),
            ("FP", "0"),
            ("FN", "0"),
            ("TN", "4"),
            ("Accuracy", "100.0%"),
            ("Precision", "100.0%"),
            ("Recall", "100.0%"),
        ]
    ):
        t3.cell(i, 0).text = a
        t3.cell(i, 1).text = b
    p('Unit suite only — not "I labelled the whole prospectus by hand."')

    h("Other live checks", 3)
    bullet("manual_test.py — passed (all 9 types)")
    bullet("generic_docx_test.py — passed (3 layouts)")
    bullet("ml_ner_test.py — passed (Rules 2 vs Hybrid 5)")

    h("What's inside the submitted redacted prospectus (counted today)", 3)
    t4 = doc.add_table(rows=6, cols=2)
    t4.style = "Table Grid"
    for i, (a, b) in enumerate(
        [
            ("Marker", "Count"),
            ("Non-empty paragraphs", "694"),
            ("@example.com emails", "38"),
            ("Example Entity … Limited", "77"),
            ("Example Avenue addresses", "27"),
            ("Synthetic +91 phones", "13"),
        ]
    ):
        t4.cell(i, 0).text = a
        t4.cell(i, 1).text = b
    p(
        "Original file isn't in git. To reprint a full engine summary: "
        'python verify_metrics.py --prospectus "PATH\\original.docx"'
    )

    h("Trade-offs", 2)
    p(
        "Rules miss bare names; hybrid helps but isn't perfect on legalese. "
        "Weird multi-line addresses still hurt. Free cloud for spaCy didn't stick, "
        "so local is the real ML path."
    )

    h("Grading pack", 2)
    bullet("Code — repo / redact_pii.py")
    bullet("Output — Red Herring Prospectus - Redacted.docx")
    bullet("This note + EVALUATION_REPORT.md")
    bullet("Recompute — python verify_metrics.py")

    doc.save(OUT)
    print(f"Wrote {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
