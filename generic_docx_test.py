"""Regression tests proving the redactor is not tied to one source document.

The script builds several unrelated DOCX files from scratch, runs the real
`redact_docx` engine, and verifies that seeded PII values are removed while
ordinary non-PII controls are retained.

Run:
    python generic_docx_test.py
"""

from __future__ import annotations

import json
import shutil
from dataclasses import dataclass
from pathlib import Path

from docx import Document

from redact_pii import redact_docx


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "examples" / "generic_regression"
REPORT = OUT_DIR / "generic_docx_report.json"
REQUIRED_TYPES = {"name", "email", "phone", "company", "address", "ssn", "card", "dob", "ip"}


@dataclass(frozen=True)
class Scenario:
    name: str
    expected_originals: tuple[str, ...]
    controls: tuple[str, ...]


def _clean() -> None:
    if OUT_DIR.exists():
        shutil.rmtree(OUT_DIR)
    OUT_DIR.mkdir(parents=True, exist_ok=True)


def _text(path: Path) -> str:
    document = Document(path)
    parts: list[str] = []
    for paragraph in document.paragraphs:
        parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    parts.append(paragraph.text)
    for section in document.sections:
        for story in (section.header, section.footer, section.first_page_header, section.first_page_footer):
            for paragraph in story.paragraphs:
                parts.append(paragraph.text)
            for table in story.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            parts.append(paragraph.text)
    return "\n".join(parts)


def build_customer_ticket(path: Path) -> Scenario:
    document = Document()
    document.add_heading("Support Ticket Export", level=1)
    document.add_paragraph("Ticket ID: TCK-2026-00045")
    document.add_paragraph("Customer Name: Elena Brooks")
    document.add_paragraph("Email: elena.brooks@example.org")
    document.add_paragraph("Phone: (415) 555-0198")
    document.add_paragraph("Billing Address: 742 Evergreen Terrace, Springfield, IL 62704")
    document.add_paragraph("DOB: 1988-04-22")
    document.add_paragraph("Login IP: 10.24.8.99")
    document.add_paragraph("Card ending token: 4111-1111-1111-1111")
    document.add_paragraph("SSN: 321-54-9876")
    document.add_paragraph("Employer: Northwind Traders LLC")
    document.add_paragraph("Control value should remain: Order 2026-00045")
    document.save(path)
    return Scenario(
        "customer_ticket",
        (
            "Elena Brooks",
            "elena.brooks@example.org",
            "(415) 555-0198",
            "742 Evergreen Terrace, Springfield, IL 62704",
            "1988-04-22",
            "10.24.8.99",
            "4111-1111-1111-1111",
            "321-54-9876",
            "Northwind Traders LLC",
        ),
        ("Order 2026-00045",),
    )


def build_hr_table(path: Path) -> Scenario:
    document = Document()
    document.sections[0].header.paragraphs[0].text = "Prepared by: Dr. Kavya Rao"
    document.sections[0].footer.paragraphs[0].text = "Helpdesk: +91 99887 77665"
    document.add_heading("Employee Onboarding Table", level=1)
    table = document.add_table(rows=6, cols=2)
    rows = [
        ("Full Name", "Marcus Hill"),
        ("E-mail", "marcus.hill@sample.net"),
        ("Residence", "Flat 8B, Sunrise Tower, MG Road, Pune - 411 001, India"),
        ("Company", "Blue River Analytics Pvt Ltd"),
        ("Birth Date", "7 Jan 1992"),
        ("IPv4", "172.16.5.20"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
    document.add_paragraph("Control value should remain: Employee code EMP-1042")
    document.save(path)
    return Scenario(
        "hr_table_header_footer",
        (
            "Dr. Kavya Rao",
            "+91 99887 77665",
            "Marcus Hill",
            "marcus.hill@sample.net",
            "Flat 8B, Sunrise Tower, MG Road, Pune - 411 001, India",
            "Blue River Analytics Pvt Ltd",
            "7 Jan 1992",
            "172.16.5.20",
        ),
        ("Employee code EMP-1042",),
    )


def build_split_runs(path: Path) -> Scenario:
    document = Document()
    document.add_heading("Run-Split Contact Note", level=1)
    paragraph = document.add_paragraph()
    paragraph.add_run("Contact Person: ")
    paragraph.add_run("Sofia").bold = True
    paragraph.add_run(" Martinez")
    paragraph = document.add_paragraph()
    paragraph.add_run("Reach at ")
    paragraph.add_run("sofia.martinez@example.com").italic = True
    document.add_paragraph("Mobile: 9876543210")
    document.add_paragraph("Mailing Address: 12 Market Street, San Francisco, CA 94105")
    document.add_paragraph("SSN: 555-66-7777")
    document.add_paragraph("DOB: 03/09/1991")
    document.add_paragraph("Company: Contoso Corporation")
    document.add_paragraph("Card: 5555 5555 5555 4444")
    document.add_paragraph("IP address: 203.0.113.45")
    document.add_paragraph("Control value should remain: Invoice 2026-00991")
    document.save(path)
    return Scenario(
        "split_runs",
        (
            "Sofia Martinez",
            "sofia.martinez@example.com",
            "9876543210",
            "12 Market Street, San Francisco, CA 94105",
            "555-66-7777",
            "03/09/1991",
            "Contoso Corporation",
            "5555 5555 5555 4444",
            "203.0.113.45",
        ),
        ("Invoice 2026-00991",),
    )


def main() -> None:
    _clean()
    builders = (build_customer_ticket, build_hr_table, build_split_runs)
    results = []
    all_passed = True
    for builder in builders:
        input_path = OUT_DIR / f"{builder.__name__.replace('build_', '')}.docx"
        scenario = builder(input_path)
        output_path = OUT_DIR / f"{scenario.name}_redacted.docx"
        summary = redact_docx(input_path, output_path)
        result_text = _text(output_path)
        remaining = [value for value in scenario.expected_originals if value in result_text]
        missing_controls = [value for value in scenario.controls if value not in result_text]
        detected_types = set(summary["redactions"])
        passed = not remaining and not missing_controls and detected_types >= (REQUIRED_TYPES & detected_types)
        if scenario.name in {"customer_ticket", "split_runs"}:
            passed = passed and REQUIRED_TYPES <= detected_types
        else:
            passed = passed and {"name", "email", "phone", "company", "address", "dob", "ip"} <= detected_types
        all_passed = all_passed and passed
        results.append(
            {
                "scenario": scenario.name,
                "passed": passed,
                "input": str(input_path),
                "output": str(output_path),
                "redactions": summary["redactions"],
                "original_values_remaining": remaining,
                "missing_controls": missing_controls,
            }
        )
    report = {"passed": all_passed, "scenarios": results}
    REPORT.write_text(json.dumps(report, indent=2), encoding="utf-8")
    print(json.dumps(report, indent=2))
    if not all_passed:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
