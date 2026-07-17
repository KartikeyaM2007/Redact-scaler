"""Create and verify a small end-to-end DOCX redaction example.

Run this from the project folder:
    python manual_test.py

It produces `examples/manual_test_input.docx`, its redacted counterpart, and
`examples/manual_test_report.json`. The fixture exercises every PII type named
in the assignment and verifies that the original strings no longer appear in
the resulting DOCX text.
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

from redact_pii import redact_docx


ROOT = Path(__file__).resolve().parent
EXAMPLES = ROOT / "examples"
INPUT = EXAMPLES / "manual_test_input.docx"
OUTPUT = EXAMPLES / "manual_test_redacted.docx"
REPORT = EXAMPLES / "manual_test_report.json"

FIXTURE = [
    ("Full name", "Contact Person: Rashi Patil"),
    ("Titled full name", "Backup contact: Mr. Rohan Dey"),
    ("Email", "E-mail: rashi.patil@example.com"),
    ("Phone", "Telephone: +91 98765 43210"),
    ("Company", "Company: Acme Financial Services Limited"),
    ("Mailing address", "Registered Office: 11/3 Example Road, Pune - 411 001, India"),
    ("SSN", "SSN: 123-45-6789"),
    ("Credit card", "Card: 4111 1111 1111 1111"),
    ("Date of birth", "Date of Birth: 12/03/1990"),
    ("IP address", "IP address: 192.168.1.10"),
]


def create_fixture(path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.8)
    section.left_margin = section.right_margin = Inches(0.85)
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(5)

    title = document.add_paragraph()
    title.style = document.styles["Heading 1"]
    title.add_run("PII Redaction Manual Test Fixture")
    document.add_paragraph(
        "This controlled sample contains one example for each PII type required by the assignment."
    )
    for label, value in FIXTURE:
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{label}: ").bold = True
        paragraph.add_run(value)
    document.add_paragraph("Control value (should remain): Offer closes on 10 December 2025.")
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def document_text(path: Path) -> str:
    document = Document(path)
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def main() -> None:
    create_fixture(INPUT)
    summary = redact_docx(INPUT, OUTPUT)
    result_text = document_text(OUTPUT)
    originals = [value for _, value in FIXTURE]
    missing = [value for value in originals if value in result_text]
    control_retained = "Offer closes on 10 December 2025." in result_text
    required_types = {"name", "email", "phone", "company", "address", "ssn", "card", "dob", "ip"}
    detected_types = set(summary["redactions"])
    report = {
        "passed": not missing and control_retained and required_types <= detected_types,
        "input": str(INPUT),
        "output": str(OUTPUT),
        "redactions": summary["redactions"],
        "types_exercised": sorted(detected_types),
        "original_values_remaining": missing,
        "non_pii_control_retained": control_retained,
    }
    REPORT.write_text(json.dumps(report, indent=2), encoding="utf-8")
    if not report["passed"]:
        raise SystemExit(json.dumps(report, indent=2))
    print(json.dumps(report, indent=2))


if __name__ == "__main__":
    main()
